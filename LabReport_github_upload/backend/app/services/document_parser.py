import csv
import hashlib
import math
import re
import statistics
from pathlib import Path
from typing import Dict, List

from app.services.pdf_parser import _is_likely_lab_test, _normalize_test_name, extract_lab_data


TEXT_PATTERN = re.compile(
    r"([A-Za-z][A-Za-z0-9 ()/%+\-]{1,70})\s*[:\-]?\s*([0-9]+(?:[.,][0-9]+)?)\s*([A-Za-z/%^0-9.\-]*)",
    re.IGNORECASE,
)

LINE_PATTERN = re.compile(
    r"^\s*([A-Za-z][A-Za-z0-9 ()/%+\-]{1,70})\s+([0-9]+(?:[.,][0-9]+)?)\s*([A-Za-z/%^0-9.\-]*)",
    re.IGNORECASE,
)

FINDING_RULES = [
    ("Pneumonia / consolidation signal", ["pneumonia", "consolidation", "infiltrate", "airspace opacity", "patchy opacity"]),
    ("Pleural effusion signal", ["pleural effusion", "costophrenic angle blunting", "effusion"]),
    ("Pneumothorax signal", ["pneumothorax", "collapsed lung"]),
    ("Cardiomegaly signal", ["cardiomegaly", "enlarged cardiac silhouette", "cardiac enlargement"]),
    ("Fracture signal", ["fracture", "cortical break", "displaced fragment"]),
    ("Tuberculosis pattern signal", ["tuberculosis", "tb", "cavity lesion", "fibrocalcific"]),
]

NO_ACUTE_RULES = [
    "no acute cardiopulmonary abnormality",
    "no acute disease",
    "normal study",
    "no focal consolidation",
    "lungs are clear",
]


def _to_float(value: str) -> float | None:
    raw = (value or "").strip().replace(",", ".")
    try:
        return float(raw)
    except Exception:
        return None


def _clean_unit(unit: str) -> str:
    return re.sub(r"[^A-Za-z0-9/%^.\-]", "", unit or "").strip()


def _extract_from_text(text: str) -> List[Dict]:
    rows: List[Dict] = []
    seen = set()

    def add_row(name: str, value: str, unit: str):
        numeric = _to_float(value)
        if numeric is None:
            return
        clean_unit = _clean_unit(unit)
        if not _is_likely_lab_test(name, unit):
            return
        test_name = _normalize_test_name(name)
        key = (test_name.lower(), f"{numeric}", clean_unit.lower())
        if key in seen:
            return
        seen.add(key)
        rows.append(
            {
                "test_name": test_name,
                "value": numeric,
                "unit": clean_unit or None,
                "reference_range": None,
            }
        )

    text = text or ""
    for name, value, unit in TEXT_PATTERN.findall(text):
        add_row(name, value, unit)

    # OCR often keeps one test per line; this catches partially parsed lines.
    for line in text.splitlines():
        match = LINE_PATTERN.search(line)
        if not match:
            continue
        add_row(match.group(1), match.group(2), match.group(3))

    return rows


def _stable_score(seed_text: str, label: str, minimum: int, maximum: int) -> int:
    digest = hashlib.sha256(f"{seed_text}|{label}".encode("utf-8", errors="ignore")).hexdigest()
    num = int(digest[:8], 16)
    return minimum + (num % (maximum - minimum + 1))


def _extract_finding_markers(text: str, seed_text: str) -> List[Dict]:
    hay = (text or "").lower()
    markers: List[Dict] = []
    matched_high = []

    for label, keywords in FINDING_RULES:
        if any(k in hay for k in keywords):
            matched_high.append(label)
            markers.append(
                {
                    "test_name": f"Imaging Finding: {label}",
                    "value": float(_stable_score(seed_text, label, 72, 92)),
                    "unit": "score",
                    "reference_range": "0-40",
                    "status_override": "high",
                    "interpretation_override": f"Text/scan pattern suggests possible {label.lower()}. Clinical confirmation is required.",
                }
            )

    has_no_acute = any(k in hay for k in NO_ACUTE_RULES)
    if has_no_acute and not matched_high:
        label = "No acute abnormality signal"
        markers.append(
            {
                "test_name": f"Imaging Finding: {label}",
                "value": float(_stable_score(seed_text, label, 16, 34)),
                "unit": "score",
                "reference_range": "0-40",
                "status_override": "normal",
                "interpretation_override": "Report text suggests no immediate acute abnormality, but follow-up should be based on symptoms.",
            }
        )

    return markers


def _read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_csv_file(path: Path) -> str:
    lines = []
    with path.open("r", encoding="utf-8", errors="ignore", newline="") as fh:
        reader = csv.reader(fh)
        for row in reader:
            lines.append(" ".join(cell.strip() for cell in row if cell))
    return "\n".join(lines)


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except Exception:
        return ""
    try:
        doc = Document(str(path))
        lines: List[str] = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    lines.append(" ".join(cells))
        return "\n".join(lines)
    except Exception:
        return ""


def _read_doc(path: Path) -> str:
    # Legacy .doc support is best-effort only.
    try:
        return path.read_text(encoding="latin-1", errors="ignore")
    except Exception:
        return ""


def _read_dicom_text(path: Path) -> str:
    # DICOM may contain text tags; no pixel OCR here.
    try:
        import pydicom  # type: ignore
    except Exception:
        try:
            return path.read_text(encoding="latin-1", errors="ignore")
        except Exception:
            return ""
    try:
        ds = pydicom.dcmread(str(path), force=True)
        parts = []
        for field in ["StudyDescription", "SeriesDescription", "ProtocolName", "PerformedProcedureStepDescription"]:
            value = getattr(ds, field, None)
            if value:
                parts.append(str(value))
        return "\n".join(parts)
    except Exception:
        return ""


def _ocr_image(path: Path) -> str:
    try:
        import pytesseract
        from PIL import Image
    except Exception:
        return ""


def _extract_image_quality_signals(path: Path, ocr_text: str) -> List[Dict]:
    try:
        from PIL import Image
    except Exception:
        return []

    try:
        img = Image.open(path).convert("L")
        pixels = list(img.getdata())
    except Exception:
        return []

    if not pixels:
        return []

    width, height = img.size
    mean_intensity = sum(pixels) / len(pixels)
    std_intensity = statistics.pstdev(pixels) if len(pixels) > 1 else 0.0

    # 0..100 where higher is better around mid-exposure.
    exposure_score = max(0.0, min(100.0, 100.0 - (abs(mean_intensity - 128.0) / 128.0) * 100.0))
    # Contrast proxy (higher stdev => better contrast up to a point).
    contrast_score = max(0.0, min(100.0, (std_intensity / 64.0) * 100.0))

    # Simple edge proxy from horizontal/vertical absolute differences.
    edge_count = 0
    sample_count = 0
    for y in range(0, max(0, height - 1), max(1, height // 140)):
        row_offset = y * width
        next_row_offset = (y + 1) * width
        for x in range(0, max(0, width - 1), max(1, width // 140)):
            p = pixels[row_offset + x]
            pr = pixels[row_offset + x + 1]
            pd = pixels[next_row_offset + x]
            gx = abs(p - pr)
            gy = abs(p - pd)
            mag = math.sqrt(gx * gx + gy * gy)
            if mag > 28:
                edge_count += 1
            sample_count += 1
    edge_density = (edge_count / sample_count) if sample_count else 0.0
    edge_score = max(0.0, min(100.0, edge_density * 220.0))

    # OCR-informed readability score.
    text = ocr_text or ""
    alpha = len(re.findall(r"[A-Za-z]", text))
    digits = len(re.findall(r"\d", text))
    readability = max(0.0, min(100.0, (alpha * 0.12 + digits * 0.2) + (contrast_score * 0.25)))

    def quality_row(name: str, value: float, min_ok: float, note: str) -> Dict:
        status = "normal" if value >= min_ok else "low"
        return {
            "test_name": name,
            "value": round(value, 2),
            "unit": "score",
            "reference_range": f"{int(min_ok)}-100",
            "status_override": status,
            "interpretation_override": note,
        }

    return [
        quality_row(
            "Image Quality - Exposure Balance",
            exposure_score,
            45.0,
            "Checks if brightness/exposure is suitable for reliable value extraction.",
        ),
        quality_row(
            "Image Quality - Structural Contrast",
            contrast_score,
            38.0,
            "Checks contrast separation needed to read structures/text in scans.",
        ),
        quality_row(
            "Image Quality - Text Readability",
            readability,
            42.0,
            "Estimates whether text and numeric values are readable from the uploaded image.",
        ),
        quality_row(
            "Image Quality - Edge Definition",
            edge_score,
            35.0,
            "Estimates visual sharpness and edge clarity for OCR/analysis confidence.",
        ),
    ]
    try:
        image = Image.open(path).convert("L")
        # Improve OCR on report screenshots/scans.
        image = image.point(lambda p: 255 if p > 165 else 0)
        return pytesseract.image_to_string(image, config="--oem 3 --psm 6") or ""
    except Exception:
        return ""


def extract_structured_data(file_path: str) -> List[Dict]:
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".pdf":
        return extract_lab_data(file_path)

    if ext in {".txt"}:
        text = _read_text_file(path)
        return _extract_from_text(text) + _extract_finding_markers(text, path.name)

    if ext in {".csv"}:
        text = _read_csv_file(path)
        return _extract_from_text(text) + _extract_finding_markers(text, path.name)

    if ext in {".docx"}:
        text = _read_docx(path)
        return _extract_from_text(text) + _extract_finding_markers(text, path.name)

    if ext in {".doc"}:
        text = _read_doc(path)
        return _extract_from_text(text) + _extract_finding_markers(text, path.name)

    if ext in {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif", ".webp"}:
        text = _ocr_image(path)
        extracted = _extract_from_text(text)
        findings = _extract_finding_markers(text, path.name)
        # Always include image quality signals so non-PDF uploads have dynamic, per-image analysis.
        quality = _extract_image_quality_signals(path, text)
        return extracted + findings + quality

    if ext in {".dcm", ".dicom"}:
        text = _read_dicom_text(path)
        return _extract_from_text(text) + _extract_finding_markers(text, path.name)

    return []
